import io
import re
import openpyxl
from pypdf import PdfReader
from docx import Document as DocxDocument

class DocumentParser:
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from a PDF file using pypdf."""
        try:
            reader = PdfReader(io.BytesIO(file_content))
            text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            return "\n".join(text)
        except Exception as e:
            print(f"Error parsing PDF: {e}")
            raise ValueError(f"Failed to parse PDF document: {str(e)}")

    @staticmethod
    def parse_pothys_excel(file_content: bytes) -> dict:
        """
        Validate structure of Pothys daily branch template,
        extract all summary metrics, employees performance, and scheme totals.
        """
        try:
            wb = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
        except Exception:
            raise ValueError("Invalid report format. Please upload the official Pothys Daily Report template.")

        # 1. Sheet name verification
        req_sheets = ["Branch Summary", "Employee Performance", "Scheme Summary"]
        for s in req_sheets:
            if s not in wb.sheetnames:
                raise ValueError("Invalid report format. Please upload the official Pothys Daily Report template.")

        ws1 = wb["Branch Summary"]
        ws2 = wb["Employee Performance"]
        ws3 = wb["Scheme Summary"]

        # 2. Structural label verification
        try:
            # Sheet 1
            if ws1["A8"].value != "Gold Sales (Amount)" or ws1["A12"].value != "Total Revenue":
                raise ValueError("Invalid report format. Please upload the official Pothys Daily Report template.")
            if ws1["C8"].value != "DigiGold Enrollments" or ws1["C12"].value != "Employees Present":
                raise ValueError("Invalid report format. Please upload the official Pothys Daily Report template.")
            if ws1["A15"].value != "Customer Complaints" or ws1["A17"].value != "Manager Remarks":
                raise ValueError("Invalid report format. Please upload the official Pothys Daily Report template.")
            
            # Sheet 2 headers in Row 1
            sh2_headers = [ws2.cell(row=1, column=c).value for c in range(1, 11)]
            expected_headers = [
                "Employee Name", "Designation", "Gold Grams Sold", "Gold Amount",
                "Silver Grams Sold", "Silver Amount", "Platinum Amount", "Diamond Amount",
                "DigiGold Enrollments", "DigiSilver Enrollments"
            ]
            if sh2_headers != expected_headers:
                raise ValueError("Invalid report format. Please upload the official Pothys Daily Report template.")
        except Exception:
            raise ValueError("Invalid report format. Please upload the official Pothys Daily Report template.")

        # 3. Extract Sheet 1
        data = {
            "summary": {
                "report_date": ws1["B4"].value,
                "branch_name": ws1["D4"].value,
                "manager_name": ws1["B5"].value,
                "sub_manager_name": ws1["D5"].value,
                "gold_sales": float(ws1["B8"].value or 0.0),
                "silver_sales": float(ws1["B9"].value or 0.0),
                "platinum_sales": float(ws1["B10"].value or 0.0),
                "diamond_sales": float(ws1["B11"].value or 0.0),
                "total_revenue": float(ws1["B12"].value or 0.0),
                "digigold_enrollments": int(ws1["D8"].value or 0),
                "digisilver_enrollments": int(ws1["D9"].value or 0),
                "employees_present": int(ws1["D12"].value or 0),
                "employees_absent": int(ws1["D13"].value or 0),
                "customer_complaints": ws1["B15"].value or "None",
                "operational_issues": ws1["B16"].value or "None",
                "remarks": ws1["B17"].value or "None",
            },
            "employees": [],
            "scheme_summary": {
                "digigold_total": int(ws3["B3"].value or 0),
                "digisilver_total": int(ws3["B4"].value or 0),
                "overall_remarks": ws3["B5"].value or "None",
            }
        }

        # Extract Sheet 2 (Employee rows)
        for r in range(2, 500):
            emp_name = ws2.cell(row=r, column=1).value
            if not emp_name:
                break
            data["employees"].append({
                "name": emp_name,
                "designation": ws2.cell(row=r, column=2).value or "Sales Executive",
                "gold_grams_sold": float(ws2.cell(row=r, column=3).value or 0.0),
                "gold_amount": float(ws2.cell(row=r, column=4).value or 0.0),
                "silver_grams_sold": float(ws2.cell(row=r, column=5).value or 0.0),
                "silver_amount": float(ws2.cell(row=r, column=6).value or 0.0),
                "platinum_amount": float(ws2.cell(row=r, column=7).value or 0.0),
                "diamond_amount": float(ws2.cell(row=r, column=8).value or 0.0),
                "digigold_enrollments": int(ws2.cell(row=r, column=9).value or 0),
                "digisilver_enrollments": int(ws2.cell(row=r, column=10).value or 0),
            })

        return data

    @staticmethod
    def extract_text_and_data_from_excel(file_content: bytes) -> tuple[dict, str]:
        """
        Extract text and structure from Excel using openpyxl.
        If it's the official Pothys template, parses full structured metrics.
        """
        try:
            # Attempt to parse Pothys custom structure first
            pothys_data = DocumentParser.parse_pothys_excel(file_content)
            sum_data = pothys_data["summary"]
            
            # Populate metrics payload matching DailyReport columns
            extracted_metrics = {
                "sales_amount": sum_data["total_revenue"],
                "attendance_count": sum_data["employees_present"],
                "target_achievement": 100.0, # default/parsed achievement
                "remarks": sum_data["remarks"],
                "issues": sum_data["operational_issues"],
                "pothys_data": pothys_data # pass full dict forward for DB insertion
            }
            
            full_text = f"Pothys Swarna Mahal Daily Report\nDate: {sum_data['report_date']}\nBranch: {sum_data['branch_name']}\nTotal Revenue: {sum_data['total_revenue']}"
            return extracted_metrics, full_text
        except ValueError as ve:
            # If specifically structural validation error, propagate it
            if "Invalid report format" in str(ve):
                raise ve
            
            # Otherwise fallback to general keyword extraction
            return {}, ""
        except Exception as e:
            print(f"Error parsing Excel: {e}")
            raise ValueError(f"Failed to parse Excel document: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from Word Document using python-docx."""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text.append(paragraph.text)
            
            # Extract table contents too
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text]
                    if row_text:
                        text.append(" | ".join(row_text))
            
            return "\n".join(text)
        except Exception as e:
            print(f"Error parsing Word Document: {e}")
            raise ValueError(f"Failed to parse Word Document: {str(e)}")

    @classmethod
    def parse_document(cls, file_content: bytes, file_name: str) -> tuple[dict, str]:
        """
        Main routing method to parse a document based on its extension.
        Returns a tuple (extracted_metrics, full_text_content).
        """
        ext = file_name.split(".")[-1].lower()
        full_text = ""
        extracted_metrics = {}

        if ext == "pdf":
            full_text = cls.extract_text_from_pdf(file_content)
        elif ext in ["xlsx", "xls"]:
            extracted_metrics, full_text = cls.extract_text_and_data_from_excel(file_content)
        elif ext in ["docx", "doc"]:
            full_text = cls.extract_text_from_docx(file_content)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        # Run heuristic regex parses on PDF/Word text if structured data is not yet found
        if "sales_amount" not in extracted_metrics:
            sales_match = re.search(r"(?:sales|revenue|collection)[:\s]*Rs\.?\s*([\d,]+(?:\.\d{2})?)", full_text, re.IGNORECASE)
            if sales_match:
                extracted_metrics["sales_amount"] = float(sales_match.group(1).replace(",", ""))
        
        if "attendance_count" not in extracted_metrics:
            att_match = re.search(r"(?:attendance|staff present|headcount)[:\s]*(\d+)", full_text, re.IGNORECASE)
            if att_match:
                extracted_metrics["attendance_count"] = int(att_match.group(1))

        if "target_achievement" not in extracted_metrics:
            target_match = re.search(r"(?:target achievement|achievement|target)[:\s]*(\d+(?:\.\d+)?)\s*%", full_text, re.IGNORECASE)
            if target_match:
                extracted_metrics["target_achievement"] = float(target_match.group(1))

        # Scan for remarks / issues
        remarks_match = re.search(r"(?:remarks|notes|comments)[:\s]*(.*)", full_text, re.IGNORECASE)
        if remarks_match:
            extracted_metrics["remarks"] = remarks_match.group(1).strip()[:1000]

        issues_match = re.search(r"(?:issues|complaints|problems|incidents)[:\s]*(.*)", full_text, re.IGNORECASE)
        if issues_match:
            extracted_metrics["issues"] = issues_match.group(1).strip()[:1000]

        return extracted_metrics, full_text

document_parser = DocumentParser()
